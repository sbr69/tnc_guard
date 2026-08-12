import io
import pdfplumber
import docx
from ..models.base import CamelModel
from ..services.html_text import extract_page_text
from ..services.http_client import safe_get

_PARSE_HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.5",
    "Sec-Fetch-Dest": "document",
    "Sec-Fetch-Mode": "navigate",
    "Sec-Fetch-Site": "none",
    "Upgrade-Insecure-Requests": "1",
}

class PageContent(CamelModel):
    page_number: int
    text: str

class SectionNode(CamelModel):
    title: str
    content: str
    level: int

class ParsedDocument(CamelModel):
    raw_text: str
    pages: list[PageContent]
    sections: list[SectionNode]

def parse_pdf(file_bytes: bytes) -> ParsedDocument:
    raw_text = ""
    pages = []
    sections = []
    
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for idx, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            pages.append(PageContent(page_number=idx + 1, text=page_text))
            raw_text += page_text + "\n"
            
    # Simple rule-based headers extraction for sections
    lines = raw_text.split("\n")
    current_section_title = "Introduction"
    current_section_content = ""
    
    for line in lines:
        stripped = line.strip()
        # If the line is short, in all caps or starts with Section/Article, treat as a potential header
        if len(stripped) < 100 and (stripped.isupper() or stripped.startswith(("SECTION", "Section", "ARTICLE", "Article", "§"))):
            if current_section_content.strip():
                sections.append(SectionNode(
                    title=current_section_title,
                    content=current_section_content.strip(),
                    level=1
                ))
            current_section_title = stripped
            current_section_content = ""
        else:
            current_section_content += line + "\n"
            
    if current_section_content.strip():
        sections.append(SectionNode(
            title=current_section_title,
            content=current_section_content.strip(),
            level=1
        ))
        
    if not raw_text.strip():
        raise Exception("No readable text found in PDF. If this is a scanned document, OCR is not supported. Please paste the text directly or upload a digital PDF/Word document.")
        
    return ParsedDocument(raw_text=raw_text, pages=pages, sections=sections)

def parse_docx(file_bytes: bytes) -> ParsedDocument:
    doc = docx.Document(io.BytesIO(file_bytes))
    raw_text = ""
    pages = []
    sections = []
    
    current_section_title = "Introduction"
    current_section_content = ""
    
    for p in doc.paragraphs:
        text = p.text
        raw_text += text + "\n"
        
        # docx has style information to detect headings easily
        if p.style.name.startswith("Heading"):
            if current_section_content.strip():
                sections.append(SectionNode(
                    title=current_section_title,
                    content=current_section_content.strip(),
                    level=int(p.style.name.replace("Heading", "") or 1)
                ))
            current_section_title = text
            current_section_content = ""
        else:
            current_section_content += text + "\n"
            
    if current_section_content.strip():
        sections.append(SectionNode(
            title=current_section_title,
            content=current_section_content.strip(),
            level=1
        ))
        
    # python-docx doesn't store native page numbers easily without rendering, so we put full text as Page 1
    pages.append(PageContent(page_number=1, text=raw_text))
    
    return ParsedDocument(raw_text=raw_text, pages=pages, sections=sections)

def parse_txt(file_text: str) -> ParsedDocument:
    sections = []
    lines = file_text.split("\n")
    current_section_title = "Document"
    current_section_content = ""
    
    for line in lines:
        stripped = line.strip()
        if len(stripped) < 80 and (stripped.isupper() or stripped.startswith(("SECTION", "Section", "ARTICLE", "Article"))):
            if current_section_content.strip():
                sections.append(SectionNode(
                    title=current_section_title,
                    content=current_section_content.strip(),
                    level=1
                ))
            current_section_title = stripped
            current_section_content = ""
        else:
            current_section_content += line + "\n"
            
    if current_section_content.strip():
        sections.append(SectionNode(
            title=current_section_title,
            content=current_section_content.strip(),
            level=1
        ))
        
    return ParsedDocument(
        raw_text=file_text,
        pages=[PageContent(page_number=1, text=file_text)],
        sections=sections
    )

def parse_url(url: str) -> ParsedDocument:
    # Reuse discovery's fetched text if available (avoids a duplicate HTTP
    # round-trip for policy URLs that discovery already fetched+validated).
    from ..services.url_cache import get as _url_cache_get
    cached_text = _url_cache_get(url)
    if cached_text:
        return parse_txt(cached_text)

    response = safe_get(url, _PARSE_HEADERS, 15)
    if response is None:
        raise Exception("Failed to fetch URL (blocked or unreachable).")
    response.raise_for_status()
    # Shared extraction with discovery so cached and freshly-fetched text are
    # byte-identical (cache short-circuit above is provably lossless).
    text = extract_page_text(response.text)
    return parse_txt(text)

def parse_document(file_bytes: bytes | None, filename: str | None, raw_text: str | None = None, url: str | None = None) -> ParsedDocument:
    if url is not None:
        return parse_url(url)
        
    if raw_text is not None:
        return parse_txt(raw_text)
    
    if not file_bytes or not filename:
        raise ValueError("Either raw_text, url, or file_bytes and filename must be provided.")
        
    ext = filename.split(".")[-1].lower()
    if ext == "pdf":
        return parse_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return parse_docx(file_bytes)
    else:
        # Default to raw text decode
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1")
        return parse_txt(text)
